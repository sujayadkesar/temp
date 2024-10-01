# import os
# from docx import Document

# # Function to extract data from text files
# def extract_data_from_file(file_path):
#     try:
#         with open(file_path, 'r', encoding='utf-8') as f:
#             return f.read()
#     except Exception as e:
#         print(f"Error reading {file_path}: {e}")
#         return ""

# # Function to generate the report
# def generate_report(input_folder, output_doc):
#     # Create a new Word document
#     doc = Document()

#     # Add Title to the document
#     doc.add_heading('Windows Forensic Investigation Notes', 0)

#     # Add sections for each artifact type
#     doc.add_heading('System Information', level=1)
    
#     # Add specific registry data placeholders (You can customize further)
#     doc.add_paragraph(f"Computer Name: {extract_data_from_file(os.path.join(input_folder, 'system.txt'))}")
#     doc.add_paragraph(f"Windows Version: {extract_data_from_file(os.path.join(input_folder, 'windows_version.txt'))}")
#     doc.add_paragraph(f"Timezone: {extract_data_from_file(os.path.join(input_folder, 'timezone.txt'))}")
#     doc.add_paragraph(f"Network Information: {extract_data_from_file(os.path.join(input_folder, 'network_info.txt'))}")
#     doc.add_paragraph(f"Shutdown time: {extract_data_from_file(os.path.join(input_folder, 'shutdown_time.txt'))}")
#     doc.add_paragraph(f"Defender settings: {extract_data_from_file(os.path.join(input_folder, 'defender_settings.txt'))}")

#     doc.add_heading('Users, Groups, and User Profiles', level=1)
#     doc.add_paragraph(f"Active accounts: {extract_data_from_file(os.path.join(input_folder, 'active_accounts.txt'))}")
#     doc.add_paragraph(f"Administrator Group Members: {extract_data_from_file(os.path.join(input_folder, 'admin_group.txt'))}")
#     doc.add_paragraph(f"User Profiles: {extract_data_from_file(os.path.join(input_folder, 'user_profiles.txt'))}")

#     doc.add_heading('User Behavior', level=1)
#     doc.add_paragraph(f"UserAssist: {extract_data_from_file(os.path.join(input_folder, 'userassist.txt'))}")
#     doc.add_paragraph(f"Recent Docs: {extract_data_from_file(os.path.join(input_folder, 'recentdocs.txt'))}")
#     doc.add_paragraph(f"Shellbags: {extract_data_from_file(os.path.join(input_folder, 'shellbags.txt'))}")

#     doc.add_heading('NTFS - File System Analysis', level=1)
#     doc.add_paragraph(f"MFT Analysis: {extract_data_from_file(os.path.join(input_folder, 'mft.txt'))}")
#     doc.add_paragraph(f"USN Journal: {extract_data_from_file(os.path.join(input_folder, 'usn_journal.txt'))}")

#     doc.add_heading('Execution Artifacts', level=1)
#     doc.add_paragraph(f"BAM Executables: {extract_data_from_file(os.path.join(input_folder, 'bam.txt'))}")
#     doc.add_paragraph(f"AppCompatCache: {extract_data_from_file(os.path.join(input_folder, 'appcompatcache.txt'))}")
#     doc.add_paragraph(f"AmCache: {extract_data_from_file(os.path.join(input_folder, 'amcache.txt'))}")

#     doc.add_heading('Persistence Mechanisms', level=1)
#     doc.add_paragraph(f"Auto-Run Keys: {extract_data_from_file(os.path.join(input_folder, 'autorun_keys.txt'))}")
#     doc.add_paragraph(f"Startup Folder: {extract_data_from_file(os.path.join(input_folder, 'startup_folder.txt'))}")
    
#     doc.add_heading('Windows Services', level=1)
#     doc.add_paragraph(f"Services: {extract_data_from_file(os.path.join(input_folder, 'services.txt'))}")
    
#     doc.add_heading('Scheduled Tasks', level=1)
#     doc.add_paragraph(f"Scheduled Tasks: {extract_data_from_file(os.path.join(input_folder, 'scheduled_tasks.txt'))}")

#     doc.add_heading('Windows Event Log Analysis', level=1)
#     doc.add_paragraph(f"Event Logs: {extract_data_from_file(os.path.join(input_folder, 'event_logs.txt'))}")
    
#     doc.add_heading('Memory Analysis', level=1)
#     doc.add_paragraph(f"Memory Artifacts: {extract_data_from_file(os.path.join(input_folder, 'memory_analysis.txt'))}")
    
#     # Save the document
#     doc.save(output_doc)
#     print(f"Report generated: {output_doc}")

# # Main function
# if __name__ == "__main__":
#     # Folder containing parsed registry artifacts
#     input_folder = input("Enter the folder path containing parsed registry artifacts: ")

#     # Output report file path
#     output_doc = "Forensic_Report.docx"

#     # Generate the report
#     generate_report(input_folder, output_doc)


import os
from docx import Document

# Function to locate a file in a directory or its subdirectories
def find_file(root_dir, file_name):
    for dirpath, _, filenames in os.walk(root_dir):
        if file_name in filenames:
            return os.path.join(dirpath, file_name)
    return None

# Function to extract data from text files
def extract_data_from_file(file_path):
    try:
        if file_path and os.path.exists(file_path):
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            return f"{os.path.basename(file_path)} not found."
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return ""

# Function to generate the report
def generate_report(input_folder, output_doc):
    doc = Document()
    doc.add_heading('Windows Forensic Investigation Notes', 0)

    # System Information Section
    doc.add_heading('System Information', level=1)
    doc.add_paragraph(f"Computer Name: {extract_data_from_file(find_file(input_folder, 'system.txt'))}")
    doc.add_paragraph(f"Windows Version: {extract_data_from_file(find_file(input_folder, 'windows_version.txt'))}")
    doc.add_paragraph(f"Timezone: {extract_data_from_file(find_file(input_folder, 'timezone.txt'))}")
    doc.add_paragraph(f"Network Information: {extract_data_from_file(find_file(input_folder, 'network_info.txt'))}")
    doc.add_paragraph(f"Shutdown time: {extract_data_from_file(find_file(input_folder, 'shutdown_time.txt'))}")
    doc.add_paragraph(f"Defender settings: {extract_data_from_file(find_file(input_folder, 'defender_settings.txt'))}")

    # Users Groups and User Profiles Section
    doc.add_heading('Users, Groups, and User Profiles', level=1)
    doc.add_paragraph(f"Active accounts: {extract_data_from_file(find_file(input_folder, 'active_accounts.txt'))}")
    doc.add_paragraph(f"Administrator Group Members: {extract_data_from_file(find_file(input_folder, 'admin_group.txt'))}")
    doc.add_paragraph(f"User Profiles: {extract_data_from_file(find_file(input_folder, 'user_profiles.txt'))}")

    # User Behavior Section
    doc.add_heading('User Behavior', level=1)
    doc.add_paragraph(f"UserAssist: {extract_data_from_file(find_file(input_folder, 'userassist.txt'))}")
    doc.add_paragraph(f"Recent Docs: {extract_data_from_file(find_file(input_folder, 'recentdocs.txt'))}")
    doc.add_paragraph(f"Shellbags: {extract_data_from_file(find_file(input_folder, 'shellbags.txt'))}")

    # NTFS - File System Analysis Section
    doc.add_heading('NTFS - File System Analysis', level=1)
    doc.add_paragraph(f"MFT Analysis: {extract_data_from_file(find_file(input_folder, 'mft.txt'))}")
    doc.add_paragraph(f"USN Journal: {extract_data_from_file(find_file(input_folder, 'usn_journal.txt'))}")

    # Execution Artifacts Section
    doc.add_heading('Execution Artifacts', level=1)
    doc.add_paragraph(f"BAM Executables: {extract_data_from_file(find_file(input_folder, 'bam.txt'))}")
    doc.add_paragraph(f"AppCompatCache: {extract_data_from_file(find_file(input_folder, 'appcompatcache.txt'))}")
    doc.add_paragraph(f"AmCache: {extract_data_from_file(find_file(input_folder, 'amcache.txt'))}")

    # Persistence Mechanisms Section
    doc.add_heading('Persistence Mechanisms', level=1)
    doc.add_paragraph(f"Auto-Run Keys: {extract_data_from_file(find_file(input_folder, 'autorun_keys.txt'))}")
    doc.add_paragraph(f"Startup Folder: {extract_data_from_file(find_file(input_folder, 'startup_folder.txt'))}")

    # Services and Scheduled Tasks Section
    doc.add_heading('Windows Services', level=1)
    doc.add_paragraph(f"Services: {extract_data_from_file(find_file(input_folder, 'services.txt'))}")

    doc.add_heading('Scheduled Tasks', level=1)
    doc.add_paragraph(f"Scheduled Tasks: {extract_data_from_file(find_file(input_folder, 'scheduled_tasks.txt'))}")

    # Windows Event Log Analysis Section
    doc.add_heading('Windows Event Log Analysis', level=1)
    doc.add_paragraph(f"Event Logs: {extract_data_from_file(find_file(input_folder, 'event_logs.txt'))}")

    # Memory Analysis Section
    doc.add_heading('Memory Analysis', level=1)
    doc.add_paragraph(f"Memory Artifacts: {extract_data_from_file(find_file(input_folder, 'memory_analysis.txt'))}")

    # Save the document
    doc.save(output_doc)
    print(f"Report generated: {output_doc}")

# Main function
if __name__ == "__main__":
    # Enter the folder path
    input_folder = input("Enter the root folder path containing parsed registry artifacts: ")
    # Output report file path
    output_doc = "Forensic_Report.docx"
    # Generate the report
    generate_report(input_folder, output_doc)
